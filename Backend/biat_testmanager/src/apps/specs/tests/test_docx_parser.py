from io import BytesIO
from types import SimpleNamespace

from django.test import SimpleTestCase

from apps.specs.services.parsers.docx_parser import DOCXSpecificationSourceParser


class DOCXSpecificationSourceParserTests(SimpleTestCase):
    def test_parser_groups_headings_bullets_and_tables_into_meaningful_records(self):
        from docx import Document

        document = Document()
        document.add_paragraph(
            "1/Create ENTREPRISE.INTR.GF.001.JOB.RX to generate invoices for RXT and RXS"
        )
        document.add_paragraph("")
        document.add_paragraph("Facture RXT")
        document.add_paragraph("• Select claims for company RXT")
        document.add_paragraph("• Generate invoice output")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Template"
        table.cell(0, 1).text = "RXT"
        table.cell(1, 0).text = "Format"
        table.cell(1, 1).text = "Excel"

        document.add_paragraph("Facture RXS")
        document.add_paragraph("• Select claims for company RXS")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Template"
        table.cell(0, 1).text = "RXS"

        document.add_paragraph("2/Update ENTREPRISE.INTR.GF.001.JOB")
        document.add_paragraph("Refresh claim list for company TR")

        payload = BytesIO()
        document.save(payload)
        payload.seek(0)
        payload.name = "sample.docx"

        source = SimpleNamespace(name="Sample import", file=payload)

        result = DOCXSpecificationSourceParser().parse(source)

        self.assertEqual(len(result.records), 3)
        self.assertEqual(result.records[0].title, "Facture RXT")
        self.assertEqual(result.records[1].title, "Facture RXS")
        self.assertEqual(result.records[2].title, "2/Update ENTREPRISE.INTR.GF.001.JOB")

        self.assertIn("Contexte: 1/Create ENTREPRISE.INTR.GF.001.JOB.RX", result.records[0].content)
        self.assertIn("- Select claims for company RXT", result.records[0].content)
        self.assertIn("Table:", result.records[0].content)
        self.assertIn("Template: RXT", result.records[0].content)

        self.assertEqual(result.records[0].section_label, "Facture RXT")
        self.assertEqual(result.records[0].external_reference, "ENTREPRISE.INTR.GF.001.JOB.RX")
        self.assertEqual(result.source_metadata["table_count"], 2)
        self.assertEqual(result.source_metadata["latest_parse_strategy"], "structured_docx")

    def test_parser_uses_word_heading_styles_when_numbering_is_missing(self):
        from docx import Document

        document = Document()
        document.add_heading("Customer portal authentication", level=1)
        document.add_heading("Successful login", level=2)
        document.add_paragraph("User opens the login page")
        document.add_paragraph("User submits valid credentials")

        payload = BytesIO()
        document.save(payload)
        payload.seek(0)
        payload.name = "headed.docx"

        source = SimpleNamespace(name="Headed import", file=payload)

        result = DOCXSpecificationSourceParser().parse(source)

        self.assertEqual(len(result.records), 1)
        self.assertEqual(result.records[0].title, "Successful login")
        self.assertEqual(result.records[0].section_label, "Successful login")
        self.assertIn("Contexte: Customer portal authentication", result.records[0].content)
